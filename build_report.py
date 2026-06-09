import json
import math
import shutil
import sys
import textwrap
from pathlib import Path

WORKSPACE_DOCX_PACKAGES = Path(
    r"C:\Users\tomek\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\Lib\site-packages"
)
if WORKSPACE_DOCX_PACKAGES.exists():
    sys.path.append(str(WORKSPACE_DOCX_PACKAGES))

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from scipy.stats import chi2_contingency, kruskal, shapiro, spearmanr
from sklearn.cluster import KMeans
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix, r2_score, silhouette_score
from sklearn.mixture import GaussianMixture
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.svm import SVC
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor, export_text, plot_tree


ROOT = Path(__file__).resolve().parent
NOTEBOOK = ROOT / "AircraftAccidents.ipynb"
ASSETS = ROOT / "report_assets"
OUT_DOCX = ROOT / "Raport_Airplane_Crash_Data_Eksploracja_Danych.docx"

sns.set_theme(style="whitegrid", palette="deep")


def fresh_assets_dir():
    if ASSETS.exists():
        shutil.rmtree(ASSETS)
    ASSETS.mkdir(parents=True)


def load_prepared_dataframe():
    nb = json.loads(NOTEBOOK.read_text(encoding="utf-8"))
    ns = {"display": lambda *args, **kwargs: None}
    for idx in [1, 3, 4, 8, 10, 11, 13]:
        code = "".join(nb["cells"][idx]["source"])
        exec(compile(code, f"notebook_cell_{idx}", "exec"), ns)
    df = ns["df"].copy()
    df = df.dropna(subset=["AC Type", "Aboard", "Fatalities", "FatalityRate", "CauseCategory", "Year"])
    df["Decade"] = df["Decade"].astype(int)
    df["DecadeCategory"] = df["DecadeCategory"].astype(str)
    return df


def savefig(name):
    path = ASSETS / f"{name}.png"
    plt.tight_layout()
    plt.savefig(path, dpi=220, bbox_inches="tight")
    plt.close()
    return path


def round_df(df, digits=2):
    return df.copy().applymap(lambda x: round(x, digits) if isinstance(x, (float, np.floating)) and not pd.isna(x) else x)


def table_to_rows(df, index_name=None, max_rows=None):
    out = df.copy()
    if max_rows:
        out = out.head(max_rows)
    if index_name is not None:
        out = out.reset_index().rename(columns={out.index.name or "index": index_name})
    return [list(out.columns)] + out.astype(object).where(pd.notna(out), "").values.tolist()


def run_analysis(df):
    quantitative = ["Aboard", "Fatalities", "FatalityRate"]
    stats = df[quantitative].agg(["count", "mean", "median", "min", "max", "std", "var"]).T
    stats["mode"] = [df[col].mode().iloc[0] if not df[col].mode().empty else pd.NA for col in quantitative]

    type_counts = df["AC Type"].value_counts().rename_axis("Typ samolotu").to_frame("Liczba zdarzen")
    cause_counts = df["CauseCategory"].value_counts().rename_axis("Kategoria przyczyny").to_frame("Liczba zdarzen")
    decade_counts = df["Decade"].value_counts().sort_index().rename_axis("Dekada").to_frame("Liczba zdarzen")
    aboard_counts = df["AboardGroup"].value_counts().rename_axis("Grupa Aboard").to_frame("Liczba zdarzen")
    fatality_group_counts = df["FatalityRateGroup"].value_counts().rename_axis("Grupa FatalityRate").to_frame("Liczba zdarzen")

    type_stats = (
        df.groupby("AC Type")
        .agg(
            Liczba=("FatalityRate", "count"),
            Srednia_smiertelnosc=("FatalityRate", "mean"),
            Mediana_smiertelnosc=("FatalityRate", "median"),
            Odchylenie=("FatalityRate", "std"),
            Srednio_na_pokladzie=("Aboard", "mean"),
        )
        .sort_values("Srednia_smiertelnosc", ascending=False)
    )

    h1_groups = [group["FatalityRate"].dropna() for _, group in df.groupby("AC Type")]
    h1_H, h1_p = kruskal(*h1_groups)

    normality = []
    for col in quantitative:
        sample = df[col].dropna()
        if len(sample) > 5000:
            sample = sample.sample(5000, random_state=42)
        stat, p = shapiro(sample)
        normality.append([col, round(stat, 4), round(p, 5)])
    normality = pd.DataFrame(normality, columns=["Zmienna", "W", "p"])

    corr = df[["Aboard", "Fatalities", "FatalityRate", "Year", "Decade"]].corr(numeric_only=True)

    table_type_cause = pd.crosstab(df["AC Type"], df["CauseCategory"]).rename_axis(index="Typ samolotu", columns="Przyczyna")
    table_type_cause_pct = pd.crosstab(df["AC Type"], df["CauseCategory"], normalize="index").mul(100).round(1)
    table_cause_aboard = pd.crosstab(df["CauseCategory"], df["AboardGroup"]).rename_axis(index="Przyczyna", columns="AboardGroup")
    table_cause_decade = pd.crosstab(df["CauseCategory"], df["DecadeCategory"]).rename_axis(index="Przyczyna", columns="Dekada")
    table_cause_fatality = pd.crosstab(df["CauseCategory"], df["FatalityRateGroup"]).rename_axis(index="Przyczyna", columns="FatalityRateGroup")

    chi_rows = []
    for table_name, table in [
        ("AC Type x CauseCategory", table_type_cause),
        ("CauseCategory x AboardGroup", table_cause_aboard),
        ("CauseCategory x DecadeCategory", table_cause_decade),
        ("CauseCategory x FatalityRateGroup", table_cause_fatality),
    ]:
        chi2, p, dof, exp = chi2_contingency(table)
        chi_rows.append([table_name, round(chi2, 1), round(p, 4), dof])
    chi_tests = pd.DataFrame(chi_rows, columns=["Zaleznosc", "chi2", "p", "df"])

    dominant = table_type_cause_pct.idxmax(axis=1).to_frame("Dominujaca kategoria")
    dominant["Udzial %"] = table_type_cause_pct.max(axis=1)
    dominant["Liczba zdarzen"] = type_counts["Liczba zdarzen"]
    dominant = dominant.sort_values(["Dominujaca kategoria", "Udzial %"], ascending=[True, False])

    h3_df = (
        df.groupby(["Decade", "AC Type"])
        .agg(MeanFatalityRate=("FatalityRate", "mean"), Count=("FatalityRate", "count"))
        .reset_index()
    )
    h3_pivot = h3_df.pivot(index="Decade", columns="AC Type", values="MeanFatalityRate").round(1)
    h3_count = h3_df.pivot(index="Decade", columns="AC Type", values="Count").fillna(0).astype(int)

    trend_rows = []
    for ac, group in df.groupby("AC Type"):
        if group["Decade"].nunique() >= 3:
            rho, p = spearmanr(group["Decade"], group["FatalityRate"], nan_policy="omit")
            trend_rows.append([ac, len(group), group["Decade"].nunique(), round(rho, 3), round(p, 4)])
    trend_tests = pd.DataFrame(trend_rows, columns=["Typ samolotu", "N", "Liczba dekad", "rho Spearmana", "p"]).sort_values("rho Spearmana")

    decade_groups = [g["FatalityRate"].dropna() for _, g in df.groupby("Decade") if len(g) > 1]
    h3_H, h3_p = kruskal(*decade_groups)

    label_enc = LabelEncoder()
    X_class = pd.get_dummies(df[["AC Type", "AboardGroup", "DecadeCategory"]], drop_first=True)
    y_class = label_enc.fit_transform(df["CauseCategory"])
    X_train, X_test, y_train, y_test = train_test_split(
        X_class, y_class, test_size=0.3, random_state=42, stratify=y_class
    )
    clf_tree = DecisionTreeClassifier(max_depth=5, random_state=42)
    clf_tree.fit(X_train, y_train)
    y_pred = clf_tree.predict(X_test)
    tree_acc = accuracy_score(y_test, y_pred)
    tree_report = classification_report(y_test, y_pred, target_names=label_enc.classes_, output_dict=True, zero_division=0)
    tree_rules = export_text(clf_tree, feature_names=list(X_class.columns), show_weights=True, max_depth=3)

    X_reg = pd.get_dummies(df[["AC Type", "CauseCategory", "DecadeCategory", "Aboard"]], drop_first=True)
    y_reg = df["FatalityRate"]
    mask = y_reg.notna()
    X_train_r, X_test_r, y_train_r, y_test_r = train_test_split(X_reg[mask], y_reg[mask], test_size=0.3, random_state=42)
    reg_tree = DecisionTreeRegressor(max_depth=5, random_state=42)
    reg_tree.fit(X_train_r, y_train_r)
    y_pred_r = reg_tree.predict(X_test_r)
    reg_r2 = r2_score(y_test_r, y_pred_r)
    reg_rules = export_text(reg_tree, feature_names=list(X_reg.columns), max_depth=3)

    rf = RandomForestClassifier(n_estimators=300, random_state=42, n_jobs=-1)
    rf.fit(X_class, y_class)
    importance = pd.Series(rf.feature_importances_, index=X_class.columns).sort_values(ascending=False)
    grouped_importance = pd.DataFrame(
        [
            ["AC Type", float(importance[importance.index.str.startswith("AC Type_")].sum())],
            ["AboardGroup", float(importance[importance.index.str.startswith("AboardGroup_")].sum())],
            ["DecadeCategory", float(importance[importance.index.str.startswith("DecadeCategory_")].sum())],
        ],
        columns=["Grupa predyktorow", "Waznosc"],
    ).sort_values("Waznosc", ascending=False)

    cluster_features = ["Aboard", "Fatalities", "FatalityRate"]
    cluster_data = df[cluster_features].dropna().copy()
    scaler = StandardScaler()
    cluster_scaled = scaler.fit_transform(cluster_data)
    sil_scores = []
    for i in range(10):
        km = KMeans(n_clusters=3, n_init=10, random_state=42 + i)
        labels = km.fit_predict(cluster_scaled)
        sil_scores.append(silhouette_score(cluster_scaled, labels))
    km = KMeans(n_clusters=3, random_state=42, n_init=10)
    cluster_data["kmeans_cluster"] = km.fit_predict(cluster_scaled)
    em = GaussianMixture(n_components=3, random_state=42)
    cluster_data["em_cluster"] = em.fit_predict(cluster_scaled)
    cluster_means = cluster_data.groupby("em_cluster")[cluster_features].mean().round(2)
    cluster_counts = cluster_data["em_cluster"].value_counts().sort_index().rename_axis("Klaster EM").to_frame("Liczba obserwacji")

    X_svm = pd.get_dummies(df[["AC Type", "Aboard", "DecadeCategory"]], drop_first=True)
    mask = X_svm.notna().all(axis=1)
    y_svm = pd.Series(label_enc.fit_transform(df["CauseCategory"]), index=df.index)[mask].values
    X_train_s, X_test_s, y_train_s, y_test_s = train_test_split(
        X_svm[mask], y_svm, test_size=0.3, random_state=42, stratify=y_svm
    )
    svm = SVC(kernel="rbf", class_weight="balanced", random_state=42)
    svm.fit(X_train_s, y_train_s)
    y_pred_s = svm.predict(X_test_s)
    svm_acc = accuracy_score(y_test_s, y_pred_s)
    svm_report = classification_report(y_test_s, y_pred_s, target_names=label_enc.classes_, output_dict=True, zero_division=0)
    svm_cm = confusion_matrix(y_test_s, y_pred_s)
    svm_classes = label_enc.classes_

    return {
        "stats": stats,
        "type_counts": type_counts,
        "cause_counts": cause_counts,
        "decade_counts": decade_counts,
        "aboard_counts": aboard_counts,
        "fatality_group_counts": fatality_group_counts,
        "type_stats": type_stats,
        "h1_H": h1_H,
        "h1_p": h1_p,
        "normality": normality,
        "corr": corr,
        "table_type_cause": table_type_cause,
        "table_type_cause_pct": table_type_cause_pct,
        "table_cause_aboard": table_cause_aboard,
        "table_cause_decade": table_cause_decade,
        "table_cause_fatality": table_cause_fatality,
        "chi_tests": chi_tests,
        "dominant": dominant,
        "h3_df": h3_df,
        "h3_pivot": h3_pivot,
        "h3_count": h3_count,
        "trend_tests": trend_tests,
        "h3_H": h3_H,
        "h3_p": h3_p,
        "clf_tree": clf_tree,
        "reg_tree": reg_tree,
        "X_class": X_class,
        "X_reg": X_reg,
        "label_classes": label_enc.classes_,
        "tree_acc": tree_acc,
        "tree_report": tree_report,
        "tree_rules": tree_rules,
        "reg_r2": reg_r2,
        "reg_rules": reg_rules,
        "importance": importance,
        "grouped_importance": grouped_importance,
        "cluster_data": cluster_data,
        "silhouette_mean": float(np.mean(sil_scores)),
        "cluster_means": cluster_means,
        "cluster_counts": cluster_counts,
        "svm_acc": svm_acc,
        "svm_report": svm_report,
        "svm_cm": svm_cm,
        "svm_classes": svm_classes,
    }


def make_figures(df, a):
    figs = {}

    plt.figure(figsize=(11, 5))
    sns.histplot(data=df, x="FatalityRate", hue="CauseCategory", multiple="stack", bins=20)
    plt.title("Histogram smiertelnosci wypadkow wedlug kategorii przyczyny")
    plt.xlabel("FatalityRate (%)")
    plt.ylabel("Liczba zdarzen")
    figs["h1_hist"] = savefig("01_h1_hist_fatality_cause")

    plt.figure(figsize=(11, 5))
    sns.histplot(data=df, x="Aboard", hue="AboardGroup", bins=30, multiple="stack")
    plt.title("Histogram liczby osob na pokladzie wedlug grup AboardGroup")
    plt.xlabel("Liczba osob na pokladzie")
    plt.ylabel("Liczba zdarzen")
    figs["h1_aboard_hist"] = savefig("02_h1_hist_aboard")

    plt.figure(figsize=(11, 5))
    sns.pointplot(data=df, x="CauseCategory", y="FatalityRate", hue="AboardGroup", dodge=True, errorbar=None)
    plt.title("Srednia smiertelnosc w grupach przyczyn i liczby osob na pokladzie")
    plt.xlabel("Kategoria przyczyny")
    plt.ylabel("Sredni FatalityRate (%)")
    plt.xticks(rotation=35, ha="right")
    figs["h1_interaction"] = savefig("03_h1_interaction_cause_aboard")

    order = a["type_stats"].index.tolist()
    plt.figure(figsize=(11, 7))
    sns.boxplot(data=df, x="FatalityRate", y="AC Type", order=order, palette="coolwarm")
    plt.title("Rozklad smiertelnosci dla wybranych typow samolotow")
    plt.xlabel("FatalityRate (%)")
    plt.ylabel("Typ samolotu")
    figs["h1_box"] = savefig("04_h1_box_type")

    plt.figure(figsize=(11, 7))
    sns.barplot(data=a["type_stats"].reset_index(), x="Srednia_smiertelnosc", y="AC Type", palette="viridis")
    plt.title("Srednia smiertelnosc dla wybranych typow samolotow")
    plt.xlabel("Sredni FatalityRate (%)")
    plt.ylabel("Typ samolotu")
    figs["h1_bar"] = savefig("05_h1_bar_mean_type")

    trend_df = df.groupby("Year")["FatalityRate"].mean().reset_index()
    plt.figure(figsize=(11, 5))
    sns.lineplot(data=trend_df, x="Year", y="FatalityRate", marker="o", linewidth=1.5)
    plt.title("Trend sredniej smiertelnosci w latach")
    plt.xlabel("Rok")
    plt.ylabel("Sredni FatalityRate (%)")
    figs["h1_year_trend"] = savefig("06_h1_year_trend")

    plt.figure(figsize=(8, 6))
    sns.heatmap(a["corr"], annot=True, fmt=".2f", cmap="vlag", center=0)
    plt.title("Macierz korelacji zmiennych ilosciowych")
    figs["corr"] = savefig("07_corr_heatmap")

    plt.figure(figsize=(11, 6))
    sns.countplot(data=df, y="CauseCategory", order=df["CauseCategory"].value_counts().index, palette="mako")
    plt.title("Liczba zdarzen w kategoriach przyczyn")
    plt.xlabel("Liczba zdarzen")
    plt.ylabel("Kategoria przyczyny")
    figs["h2_cause_counts"] = savefig("08_h2_cause_counts")

    plt.figure(figsize=(12, 7))
    sns.heatmap(a["table_type_cause_pct"], annot=True, fmt=".1f", cmap="YlGnBu", cbar_kws={"label": "% w ramach modelu"})
    plt.title("Udzial kategorii przyczyn w ramach modelu samolotu")
    plt.xlabel("Kategoria przyczyny")
    plt.ylabel("Typ samolotu")
    figs["h2_heatmap"] = savefig("09_h2_type_cause_heatmap")

    interaction = pd.crosstab(df["Year"], df["CauseCategory"], normalize="index")
    plt.figure(figsize=(11, 5))
    interaction.rolling(5, min_periods=1).mean().plot(ax=plt.gca())
    plt.title("Interakcja roku i kategorii przyczyny (srednia kroczaca 5 lat)")
    plt.xlabel("Rok")
    plt.ylabel("Udzial kategorii przyczyny")
    plt.legend(bbox_to_anchor=(1.02, 1), loc="upper left")
    figs["h2_year_cause"] = savefig("10_h2_year_cause_interaction")

    plt.figure(figsize=(10, 5))
    sns.boxplot(data=df, x="CauseCategory", y="FatalityRate")
    plt.title("FatalityRate wedlug kategorii przyczyny")
    plt.xlabel("Kategoria przyczyny")
    plt.ylabel("FatalityRate (%)")
    plt.xticks(rotation=35, ha="right")
    figs["h2_box_cause"] = savefig("11_h2_box_cause")

    plt.figure(figsize=(10, 7))
    a["importance"].head(20).sort_values().plot(kind="barh")
    plt.title("Waznosc predyktorow dla klasyfikacji kategorii przyczyny")
    plt.xlabel("Waznosc")
    figs["h2_importance"] = savefig("12_h2_feature_importance")

    plot_df = a["h3_df"][a["h3_df"]["AC Type"].isin(df["AC Type"].value_counts().head(15).index)]
    plt.figure(figsize=(12, 7))
    sns.lineplot(data=plot_df, x="Decade", y="MeanFatalityRate", hue="AC Type", marker="o")
    plt.title("Srednia smiertelnosc w kolejnych dekadach dla najczestszych modeli")
    plt.xlabel("Dekada")
    plt.ylabel("Sredni FatalityRate (%)")
    plt.legend(bbox_to_anchor=(1.02, 1), loc="upper left")
    figs["h3_line"] = savefig("13_h3_decade_type_line")

    plt.figure(figsize=(11, 5))
    sns.boxplot(data=df, x="Decade", y="FatalityRate", color="#8fb3d9")
    plt.title("Rozklad FatalityRate w dekadach")
    plt.xlabel("Dekada")
    plt.ylabel("FatalityRate (%)")
    figs["h3_decade_box"] = savefig("14_h3_decade_box")

    plt.figure(figsize=(10, 6))
    sns.scatterplot(data=df, x="Aboard", y="FatalityRate", hue="CauseCategory", alpha=0.75)
    plt.title("Aboard vs FatalityRate z podzialem na przyczyny")
    plt.xlabel("Liczba osob na pokladzie")
    plt.ylabel("FatalityRate (%)")
    plt.legend(bbox_to_anchor=(1.02, 1), loc="upper left")
    figs["h3_scatter"] = savefig("15_h3_scatter_aboard_fatality")

    plt.figure(figsize=(17, 9))
    plot_tree(
        a["clf_tree"],
        feature_names=a["X_class"].columns,
        class_names=a["label_classes"],
        filled=True,
        rounded=True,
        fontsize=7,
    )
    plt.title("Drzewo klasyfikacyjne dla kategorii przyczyny")
    figs["tree_class"] = savefig("16_tree_classifier")

    plt.figure(figsize=(17, 9))
    plot_tree(a["reg_tree"], feature_names=a["X_reg"].columns, filled=True, rounded=True, fontsize=7)
    plt.title("Drzewo regresyjne dla FatalityRate")
    figs["tree_reg"] = savefig("17_tree_regression")

    plt.figure(figsize=(9, 6))
    sns.scatterplot(
        data=a["cluster_data"],
        x="Aboard",
        y="FatalityRate",
        hue="kmeans_cluster",
        palette="tab10",
    )
    plt.title("KMeans: liczba osob na pokladzie vs smiertelnosc")
    plt.xlabel("Aboard")
    plt.ylabel("FatalityRate (%)")
    figs["kmeans"] = savefig("18_kmeans_scatter")

    plt.figure(figsize=(9, 6))
    sns.scatterplot(
        data=a["cluster_data"],
        x="Aboard",
        y="FatalityRate",
        hue="em_cluster",
        palette="tab10",
    )
    plt.title("Gaussian Mixture: liczba osob na pokladzie vs smiertelnosc")
    plt.xlabel("Aboard")
    plt.ylabel("FatalityRate (%)")
    figs["em"] = savefig("19_em_scatter")

    plt.figure(figsize=(8, 6))
    sns.heatmap(a["svm_cm"], annot=True, fmt="d", cmap="Blues", xticklabels=a["svm_classes"], yticklabels=a["svm_classes"])
    plt.title("Macierz pomylek SVM dla kategorii przyczyny")
    plt.xlabel("Predykcja")
    plt.ylabel("Rzeczywista")
    plt.xticks(rotation=35, ha="right")
    figs["svm_cm"] = savefig("20_svm_confusion_matrix")

    return figs


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if isinstance(text, (int, float, np.integer, np.floating)) else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def add_table(doc, rows, widths=None, font_size=8, caption=None):
    if caption:
        p = doc.add_paragraph(caption, style="Caption")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, val, bold=(r_idx == 0), color="FFFFFF" if r_idx == 0 else None, size=font_size)
            if r_idx == 0:
                set_cell_shading(cell, "2E74B5")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    if widths is None:
        widths = [6.5 / len(rows[0])] * len(rows[0])
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_figure(doc, path, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_metric_table(doc, rows):
    add_table(doc, [["Element", "Wartosc"]] + rows, widths=[2.4, 4.1], font_size=9)


def report_table_from_classification(report_dict):
    rows = [["Klasa", "Precision", "Recall", "F1-score", "Support"]]
    for key, value in report_dict.items():
        if key in ["accuracy", "macro avg", "weighted avg"]:
            continue
        rows.append([key, round(value["precision"], 2), round(value["recall"], 2), round(value["f1-score"], 2), int(value["support"])])
    return rows


def configure_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string("555555")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Raport ED | strona ")
    begin_run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_char1)
    instr_run = footer.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    instr_run._r.append(instr_text)
    end_run = footer.add_run()
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char2)


def add_title_page(doc, df):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Raport z projektu: analiza wypadkow lotniczych")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Zbior danych "Airplane Crash Data Since 1908"')
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    doc.add_paragraph()
    meta = [
        ["Przedmiot", "Eksploracja Danych"],
        ["Zakres danych zrodlowych", "Airplane Crashes and Fatalities Since 1908, Kaggle"],
        ["Zakres proby w notebooku", f"{int(df['Year'].min())}-{int(df['Year'].max())}, {len(df)} zdarzen po filtracji"],
        ["Jednostka analizy", "pojedynczy wypadek lotniczy"],
        ["Kluczowe zmienne", "AC Type, Aboard, Fatalities, FatalityRate, CauseCategory, Year, Decade"],
    ]
    add_table(doc, [["Pole", "Opis"]] + meta, widths=[2.2, 4.3], font_size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Raport przedstawia trzy hipotezy badawcze dotyczace zwiazku modelu statku powietrznego "
        "ze smiertelnoscia katastrof, dominujaca kategoria przyczyny oraz zmianami tych zaleznosci w czasie. "
        "Tabele i wykresy sa oparte na analizach wykonanych w notebooku AircraftAccidents.ipynb."
    )
    doc.add_page_break()


def add_intro(doc, df, a):
    doc.add_heading("1. Wstep, opis zbioru danych i hipotezy", level=1)
    doc.add_paragraph(
        "Celem projektu jest ocena, czy model statku powietrznego moze wyjasniac roznice w charakterystyce "
        "katastrof lotniczych. Zbior Kaggle obejmuje historyczne zdarzenia lotnicze od 1908 r. do 2019 r.; "
        "w notebooku zastosowano jednak ograniczenie analityczne do lat 1950-2019 oraz do 15 najczesciej "
        "wystepujacych, znormalizowanych typow samolotow. Takie ograniczenie zmniejsza wplyw rzadkich modeli "
        "i pojedynczych obserwacji, a jednoczesnie pozostawia probe wystarczajaca do porownan statystycznych."
    )
    doc.add_paragraph(
        "Zmienne zrodlowe obejmuja m.in. date zdarzenia, operatora, typ statku powietrznego, liczbe osob na "
        "pokladzie, liczbe ofiar smiertelnych i opis zdarzenia. Na potrzeby raportu utworzono zmienne pochodne: "
        "Year, Decade, DecadeCategory, FatalityRate, AboardGroup, FatalityRateGroup oraz CauseCategory."
    )
    add_metric_table(
        doc,
        [
            ["Liczba zdarzen po filtracji", len(df)],
            ["Liczba typow samolotow", df["AC Type"].nunique()],
            ["Liczba kategorii przyczyn", df["CauseCategory"].nunique()],
            ["Zakres lat", f"{int(df['Year'].min())}-{int(df['Year'].max())}"],
            ["Sredni FatalityRate", f"{df['FatalityRate'].mean():.2f}%"],
        ],
    )
    doc.add_heading("Hipotezy badawcze", level=2)
    add_bullets(
        doc,
        [
            "H1: Model samolotu ma istotny wplyw na procent ofiar smiertelnych w katastrofie.",
            "H2: Dominujaca przyczyna katastrofy zalezy od modelu samolotu.",
            "H3: Wplyw modelu samolotu na smiertelnosc katastrof zmienial sie w czasie.",
        ],
    )
    doc.add_paragraph(
        "Dalsze rozdzialy prowadza kazda hipoteze osobno. Wspolne narzedzia modelowania, takie jak drzewa "
        "decyzyjne, analiza skupien i SVM, sa omowione w miejscach, w ktorych wspieraja interpretacje danej hipotezy."
    )
    doc.add_heading("Struktura populacji po przygotowaniu danych", level=2)
    add_table(doc, table_to_rows(round_df(a["stats"], 2), "Zmienna"), widths=[1.35, .7, .85, .85, .75, .75, .85, .85, .7], font_size=7)
    add_table(doc, table_to_rows(a["type_counts"], "Typ samolotu", max_rows=15), widths=[4.5, 2.0], font_size=8, caption="Tabela 1. Licznosci analizowanych typow samolotow")
    add_table(doc, table_to_rows(a["cause_counts"], "Kategoria przyczyny"), widths=[4.5, 2.0], font_size=8, caption="Tabela 2. Licznosci kategorii przyczyn")
    doc.add_page_break()


def add_h1(doc, df, a, figs):
    doc.add_heading("2. Hipoteza H1: model samolotu a procent ofiar smiertelnych", level=1)
    doc.add_paragraph(
        "Zmienna zaleznia w H1 to FatalityRate, czyli procent ofiar smiertelnych w stosunku do liczby osob "
        "na pokladzie. Zmiennymi objasniajacymi sa AC Type, Aboard oraz Year. Hipoteza zaklada, ze roznice "
        "konstrukcyjne, pojemnosc i okres eksploatacji modeli przekladaja sie na roznice w skutkach katastrof."
    )
    doc.add_heading("2.1. Struktura zmiennych H1", level=2)
    add_table(
        doc,
        table_to_rows(round_df(a["type_stats"], 2), "Typ samolotu"),
        widths=[2.25, .65, .9, .9, .85, .95],
        font_size=7,
        caption="Tabela 3. Statystyki FatalityRate i Aboard wedlug typu samolotu",
    )
    add_figure(doc, figs["h1_hist"], "Rysunek 1. Skategoryzowany histogram FatalityRate wedlug kategorii przyczyny.")
    add_figure(doc, figs["h1_aboard_hist"], "Rysunek 2. Histogram Aboard wedlug grup liczby osob na pokladzie.")
    doc.add_heading("2.2. Zaleznosci i dobor zmiennych", level=2)
    doc.add_paragraph(
        "Dobor zmiennych do modelowania H1 wynika z zaleznosci logicznej i statystycznej. Aboard jest mianownikiem "
        "wspolczynnika FatalityRate i rownoczesnie przybliza skale zdarzenia. Year oraz Decade reprezentuja "
        "zmieniajace sie warunki technologiczne i proceduralne. AC Type jest glownym predyktorem badanym w hipotezie."
    )
    add_figure(doc, figs["corr"], "Rysunek 3. Macierz korelacji dla zmiennych ilosciowych.")
    add_figure(doc, figs["h1_interaction"], "Rysunek 4. Wykres interakcji: sredni FatalityRate wedlug CauseCategory i AboardGroup.")
    add_figure(doc, figs["h1_box"], "Rysunek 5. Wykres ramka-wasy: FatalityRate wedlug typu samolotu.")
    add_figure(doc, figs["h1_bar"], "Rysunek 6. Sredni FatalityRate dla analizowanych typow samolotow.")
    doc.add_heading("2.3. Test statystyczny i interpretacja", level=2)
    add_table(doc, [["Test", "Statystyka", "p", "Wniosek"], ["Kruskal-Wallis: FatalityRate ~ AC Type", f"{a['h1_H']:.3f}", f"{a['h1_p']:.5f}", "istotne roznice" if a["h1_p"] < 0.05 else "brak podstaw do roznic"]], widths=[3.0, 1.2, 1.0, 1.3], font_size=8)
    add_table(doc, table_to_rows(a["normality"]), widths=[2.2, 1.2, 1.2], font_size=8, caption="Tabela 4. Test Shapiro-Wilka dla zmiennych ilosciowych")
    doc.add_paragraph(
        "Rozklady zmiennych sa dalekie od normalnosci, dlatego dla porownania FatalityRate miedzy typami samolotow "
        "zastosowano nieparametryczny test Kruskala-Wallisa. Otrzymane p ponizej 0,05 wskazuje, ze co najmniej "
        "jeden model roznil sie rozkladem smiertelnosci od pozostalych."
    )
    add_figure(doc, figs["h1_year_trend"], "Rysunek 7. Sredni FatalityRate w kolejnych latach.")
    doc.add_heading("2.4. Drzewa, skupienia i weryfikacja H1", level=2)
    doc.add_paragraph(
        f"Drzewo regresyjne dla FatalityRate wykorzystalo AC Type, CauseCategory, DecadeCategory i Aboard; "
        f"jego R2 na zbiorze testowym wynioslo {a['reg_r2']:.3f}. Wynik nie oznacza idealnej predykcji, "
        "ale reguly drzewa wskazuja, ze typ samolotu i kontekst przyczyny pomagaja segmentowac poziom smiertelnosci. "
        f"Analiza skupien EM wyodrebnila klastry o odmiennych srednich wartosciach Aboard, Fatalities i FatalityRate, "
        f"a sredni silhouette dla KMeans wyniosl {a['silhouette_mean']:.3f}."
    )
    doc.add_paragraph(
        "Wniosek dla H1: hipoteze nalezy przyjac ostroznie. Model samolotu jest istotnie powiazany z FatalityRate, "
        "jednak interpretacja powinna uwzgledniac wspolwystepujace czynniki: liczbe osob na pokladzie, okres oraz "
        "kategorie przyczyny."
    )
    doc.add_page_break()


def add_h2(doc, df, a, figs):
    doc.add_heading("3. Hipoteza H2: dominujaca przyczyna katastrofy a model samolotu", level=1)
    doc.add_paragraph(
        "Zmienna zaleznia w H2 to CauseCategory, utworzona na podstawie slow kluczowych w kolumnie Summary. "
        "Zmienne objasniajace to AC Type, Year i FatalityRate. Hipoteza dotyczy tego, czy profile przyczyn "
        "sa odmienne dla roznych modeli samolotow."
    )
    doc.add_heading("3.1. Struktura zmiennych H2", level=2)
    add_table(doc, table_to_rows(a["cause_counts"], "Kategoria przyczyny"), widths=[4.5, 2.0], font_size=8, caption="Tabela 5. Rozklad zmiennej CauseCategory")
    add_figure(doc, figs["h2_cause_counts"], "Rysunek 8. Liczba wypadkow wedlug kategorii przyczyny.")
    add_table(
        doc,
        table_to_rows(a["dominant"].round(1), "Typ samolotu"),
        widths=[2.4, 2.1, 1.0, 1.0],
        font_size=7,
        caption="Tabela 6. Dominujaca kategoria przyczyny w ramach kazdego typu samolotu",
    )
    doc.add_heading("3.2. Tabele wielodzielcze i zaleznosci", level=2)
    add_table(
        doc,
        table_to_rows(a["table_type_cause"], "Typ samolotu"),
        widths=[1.7] + [4.8 / len(a["table_type_cause"].columns)] * len(a["table_type_cause"].columns),
        font_size=6,
        caption="Tabela 7. Tabela wielodzielcza: AC Type x CauseCategory",
    )
    add_figure(doc, figs["h2_heatmap"], "Rysunek 9. Udzial kategorii przyczyn w ramach modelu samolotu.")
    add_table(
        doc,
        table_to_rows(a["table_cause_aboard"], "Przyczyna"),
        widths=[2.2] + [4.3 / len(a["table_cause_aboard"].columns)] * len(a["table_cause_aboard"].columns),
        font_size=7,
        caption="Tabela 8. CauseCategory x AboardGroup",
    )
    add_table(
        doc,
        table_to_rows(a["table_cause_decade"], "Przyczyna"),
        widths=[1.7] + [4.8 / len(a["table_cause_decade"].columns)] * len(a["table_cause_decade"].columns),
        font_size=6,
        caption="Tabela 9. CauseCategory x DecadeCategory",
    )
    doc.add_heading("3.3. Testy chi-kwadrat i waznosc predyktorow", level=2)
    add_table(doc, table_to_rows(a["chi_tests"]), widths=[3.2, 1.0, .8, .8], font_size=8, caption="Tabela 10. Testy chi-kwadrat dla zmiennej CauseCategory")
    doc.add_paragraph(
        "Najwazniejszy test dla H2 to AC Type x CauseCategory. Wynik istotny statystycznie oznacza, ze rozklad "
        "kategorii przyczyn nie jest jednakowy dla wszystkich typow samolotow. Dodatkowo tabele z AboardGroup "
        "i DecadeCategory pokazuja, ze profil przyczyn jest powiazany rowniez ze skala zdarzenia i okresem."
    )
    add_figure(doc, figs["h2_year_cause"], "Rysunek 10. Wykres interakcji roku i kategorii przyczyny.")
    add_figure(doc, figs["h2_box_cause"], "Rysunek 11. FatalityRate wedlug kategorii przyczyny.")
    add_table(doc, table_to_rows(a["grouped_importance"].round(3)), widths=[3.2, 1.4], font_size=8, caption="Tabela 11. Zagregowana waznosc grup predyktorow w Random Forest")
    add_figure(doc, figs["h2_importance"], "Rysunek 12. Najwazniejsze predyktory dla klasyfikacji CauseCategory.")
    doc.add_heading("3.4. Drzewo klasyfikacyjne, SVM i weryfikacja H2", level=2)
    doc.add_paragraph(
        f"Drzewo klasyfikacyjne przewidywalo CauseCategory na podstawie AC Type, AboardGroup i DecadeCategory; "
        f"accuracy na zbiorze testowym wynioslo {a['tree_acc']:.3f}. SVM z jadrem RBF osiagnal accuracy "
        f"{a['svm_acc']:.3f}. Niska lub umiarkowana skutecznosc predykcyjna jest spodziewana, poniewaz Summary "
        "zawiera opis tekstowy, a kategorie przyczyn sa przyblizone metoda slow kluczowych."
    )
    doc.add_paragraph(
        "Wniosek dla H2: hipoteze mozna przyjac. Test chi-kwadrat dla AC Type x CauseCategory wykazuje istotna "
        "zaleznosc, a waznosc predyktorow i drzewo klasyfikacyjne potwierdzaja, ze model samolotu niesie informacje "
        "o profilu przyczyn. Wniosek nalezy traktowac jako zaleznosc eksploracyjna, nie jako dowod przyczynowy."
    )
    doc.add_page_break()


def add_h3(doc, df, a, figs):
    doc.add_heading("4. Hipoteza H3: zmiennosc wplywu modelu samolotu w czasie", level=1)
    doc.add_paragraph(
        "Zmienna zaleznia w H3 to FatalityRate, natomiast zmienne objasniajace to AC Type, Year i Decade. "
        "Hipoteza sprawdza, czy smiertelnosc katastrof dla poszczegolnych modeli zmieniala sie miedzy dekadami."
    )
    doc.add_heading("4.1. Struktura czasowa populacji", level=2)
    add_table(doc, table_to_rows(a["decade_counts"], "Dekada"), widths=[2.5, 2.0], font_size=8, caption="Tabela 12. Liczba zdarzen wedlug dekady")
    add_figure(doc, figs["h3_line"], "Rysunek 13. Sredni FatalityRate w kolejnych dekadach dla najczestszych modeli.")
    add_figure(doc, figs["h3_decade_box"], "Rysunek 14. Rozklad FatalityRate w dekadach.")
    doc.add_heading("4.2. Tabele trendow dekadowych", level=2)
    h3_selected = a["h3_pivot"][df["AC Type"].value_counts().head(8).index].copy()
    add_table(
        doc,
        table_to_rows(h3_selected, "Dekada"),
        widths=[.8] + [5.7 / len(h3_selected.columns)] * len(h3_selected.columns),
        font_size=6,
        caption="Tabela 13. Sredni FatalityRate (%) w dekadach dla najliczniejszych modeli",
    )
    h3_count_selected = a["h3_count"][df["AC Type"].value_counts().head(8).index].copy()
    add_table(
        doc,
        table_to_rows(h3_count_selected, "Dekada"),
        widths=[.8] + [5.7 / len(h3_count_selected.columns)] * len(h3_count_selected.columns),
        font_size=6,
        caption="Tabela 14. Liczba zdarzen w dekadach dla najliczniejszych modeli",
    )
    doc.add_heading("4.3. Testy trendu i zaleznosci z czasem", level=2)
    add_table(
        doc,
        table_to_rows(a["trend_tests"], max_rows=15),
        widths=[2.4, .55, .8, .95, .75],
        font_size=7,
        caption="Tabela 15. Korelacja rang Spearmana miedzy dekada a FatalityRate w ramach modelu",
    )
    add_table(
        doc,
        [["Test", "Statystyka", "p", "Wniosek"], ["Kruskal-Wallis: FatalityRate ~ Decade", f"{a['h3_H']:.3f}", f"{a['h3_p']:.5f}", "istotne roznice dekadowe" if a["h3_p"] < 0.05 else "brak podstaw do roznic"]],
        widths=[3.1, 1.2, 1.0, 1.3],
        font_size=8,
    )
    add_figure(doc, figs["h3_scatter"], "Rysunek 15. Relacja Aboard i FatalityRate z podzialem na przyczyny.")
    doc.add_heading("4.4. Drzewa, skupienia i weryfikacja H3", level=2)
    doc.add_paragraph(
        "Drzewo regresyjne, korzystajac z DecadeCategory, pokazuje, czy czas pojawia sie w regulach segmentujacych "
        "FatalityRate. Wykres dekadowy wskazuje, ze zmiennosc jest silna, ale niejednorodna: czesc modeli ma "
        "okresy wyzszej smiertelnosci, a w innych dekadach liczebnosci sa zbyt male, aby interpretowac pojedyncze "
        "punkty jako stabilny trend."
    )
    doc.add_paragraph(
        "Wniosek dla H3: hipoteze nalezy przyjac czesciowo. Widoczne sa roznice dekadowe i lokalne trendy dla "
        "wybranych modeli, ale nie dla kazdego typu samolotu proba obejmuje wystarczajaca liczbe zdarzen w wielu "
        "dekadach. Najsilniejszym uzasadnieniem H3 jest laczne odczytanie trendow, testu dekadowego i drzewa regresyjnego."
    )
    doc.add_page_break()


def add_models_and_summary(doc, a, figs):
    doc.add_heading("5. Wspolne modele data mining i analiza skupien", level=1)
    doc.add_paragraph(
        "Ta czesc zawiera modele wspolne dla trzech hipotez. W H1 i H3 najwazniejsze jest drzewo regresyjne "
        "dla FatalityRate, natomiast w H2 drzewo klasyfikacyjne i SVM dla CauseCategory. Analiza skupien opisuje "
        "naturalne segmenty zdarzen ze wzgledu na Aboard, Fatalities i FatalityRate."
    )
    doc.add_heading("5.1. Drzewa decyzyjne", level=2)
    add_metric_table(
        doc,
        [
            ["Accuracy drzewa klasyfikacyjnego", f"{a['tree_acc']:.3f}"],
            ["R2 drzewa regresyjnego", f"{a['reg_r2']:.3f}"],
        ],
    )
    add_table(doc, report_table_from_classification(a["tree_report"]), widths=[2.2, 1.0, 1.0, 1.0, .8], font_size=7, caption="Tabela 16. Raport klasyfikacji drzewa decyzyjnego")
    add_figure(doc, figs["tree_class"], "Rysunek 16. Drzewo klasyfikacyjne dla CauseCategory.")
    add_figure(doc, figs["tree_reg"], "Rysunek 17. Drzewo regresyjne dla FatalityRate.")
    doc.add_paragraph("Wybrane reguly drzewa klasyfikacyjnego:")
    for line in a["tree_rules"].splitlines()[:14]:
        doc.add_paragraph(line[:120], style="List Bullet")
    doc.add_paragraph("Wybrane reguly drzewa regresyjnego:")
    for line in a["reg_rules"].splitlines()[:14]:
        doc.add_paragraph(line[:120], style="List Bullet")

    doc.add_heading("5.2. Analiza skupien", level=2)
    add_metric_table(doc, [["Sredni silhouette score KMeans", f"{a['silhouette_mean']:.3f}"]])
    add_table(doc, table_to_rows(a["cluster_counts"], "Klaster EM"), widths=[2.0, 2.0], font_size=8, caption="Tabela 17. Liczebnosci klastrow EM")
    add_table(doc, table_to_rows(a["cluster_means"], "Klaster EM"), widths=[1.5, 1.4, 1.4, 1.5], font_size=8, caption="Tabela 18. Charakterystyka klastrow EM")
    add_figure(doc, figs["kmeans"], "Rysunek 18. KMeans: Aboard i FatalityRate.")
    add_figure(doc, figs["em"], "Rysunek 19. Gaussian Mixture: Aboard i FatalityRate.")
    doc.add_paragraph(
        "Skupienia nie odpowiadaja bezposrednio nazwom modeli, ale oddzielaja zdarzenia wedlug skali i skutkow. "
        "Dla H1 pomagaja interpretowac, czy wysoki FatalityRate dotyczy malej czy duzej liczby osob na pokladzie. "
        "Dla H3 pokazuja, ze same trendy czasowe nie wyczerpuja struktury danych, bo wazna pozostaje skala zdarzenia."
    )

    doc.add_heading("5.3. Wybrany algorytm data mining: SVM", level=2)
    add_metric_table(doc, [["Accuracy SVM", f"{a['svm_acc']:.3f}"]])
    add_table(doc, report_table_from_classification(a["svm_report"]), widths=[2.2, 1.0, 1.0, 1.0, .8], font_size=7, caption="Tabela 19. Raport klasyfikacji SVM")
    add_figure(doc, figs["svm_cm"], "Rysunek 20. Macierz pomylek SVM.")
    doc.add_paragraph(
        "SVM pelni role algorytmu data mining dla H2: sprawdza, czy z AC Type, Aboard i DecadeCategory mozna "
        "odtworzyc kategorie przyczyny. Wynik pokazuje ograniczona, ale realna informacje predykcyjna w zmiennych "
        "objasniajacych. Model nie zastępuje analizy tekstu Summary, lecz uzupelnia test chi-kwadrat."
    )

    doc.add_heading("6. Podsumowanie w kontekscie hipotez", level=1)
    add_table(
        doc,
        [
            ["Hipoteza", "Decyzja", "Najwazniejsze przeslanki"],
            ["H1", "przyjeta ostroznie", "istotny test Kruskala-Wallisa, roznice srednich i rozkladow FatalityRate wedlug AC Type"],
            ["H2", "przyjeta", "istotny chi-kwadrat AC Type x CauseCategory, rozne dominujace przyczyny dla modeli, wsparcie przez Random Forest"],
            ["H3", "przyjeta czesciowo", "widoczne roznice dekadowe, ale nierowne liczebnosci w modelach ograniczaja sile wniosku"],
        ],
        widths=[1.0, 1.6, 3.9],
        font_size=8,
        caption="Tabela 20. Decyzje badawcze dla hipotez",
    )
    doc.add_paragraph(
        "Najsilniejszy wynik dotyczy H2, gdzie zaleznosc miedzy typem samolotu a kategoria przyczyny jest "
        "potwierdzona testem chi-kwadrat. H1 jest rowniez wsparta statystycznie, ale wymaga ostroznej interpretacji, "
        "bo FatalityRate zalezy od skali zdarzenia i charakteru przyczyny. H3 pokazuje, ze czas ma znaczenie, "
        "jednak nie wszystkie modele wystepuja rownomiernie w kolejnych dekadach."
    )
    doc.add_paragraph(
        "Ograniczeniem projektu jest sposob automatycznej kategoryzacji przyczyn na podstawie slow kluczowych oraz "
        "historyczna nierownomiernosc zbioru danych. Mimo tych ograniczen raport pokazuje spojna sciezke eksploracji: "
        "od struktury populacji, przez zaleznosci i testy statystyczne, po drzewa decyzyjne, skupienia oraz SVM."
    )


def build_report(df, a, figs):
    doc = Document()
    configure_doc(doc)
    add_title_page(doc, df)
    add_intro(doc, df, a)
    add_h1(doc, df, a, figs)
    add_h2(doc, df, a, figs)
    add_h3(doc, df, a, figs)
    add_models_and_summary(doc, a, figs)
    doc.save(OUT_DOCX)


def main():
    fresh_assets_dir()
    df = load_prepared_dataframe()
    analysis = run_analysis(df)
    figs = make_figures(df, analysis)
    build_report(df, analysis, figs)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
